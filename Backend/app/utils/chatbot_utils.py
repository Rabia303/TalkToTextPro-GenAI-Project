# import whisper
# from deep_translator import GoogleTranslator
# import tempfile
# import os
# from fpdf import FPDF
# from docx import Document
# import google.generativeai as genai
# import PyPDF2
# import docx as docx_module
# from moviepy.video.io.VideoFileClip import VideoFileClip
# from datetime import datetime
# import logging

# logger = logging.getLogger(__name__)

# # Global variables for models
# whisper_model = None
# gemini_configured = False

# def init_chatbot_models(app):
#     global whisper_model, gemini_configured
#     try:
#         whisper_model = whisper.load_model("base")
#         app.logger.info("Whisper model loaded successfully")
#     except Exception as e:
#         app.logger.error(f"Failed to load Whisper model: {str(e)}")
#         whisper_model = None

#     try:
#         genai.configure(api_key=app.config['GEMINI_API_KEY'])
#         gemini_configured = True
#         app.logger.info("Gemini configured successfully")
#     except Exception as e:
#         app.logger.error(f"Failed to configure Gemini: {str(e)}")
#         gemini_configured = False

# def sanitize_text(text):
#     """Replace problematic Unicode characters with ASCII equivalents for FPDF/DOCX."""
#     if not text:
#         return ""
#     replacements = {
#         '‘': "'", '’': "'",
#         '“': '"', '”': '"',
#         '—': '-', '–': '-',
#         '…': '...',
#         '•': '-',    # replace bullet with dash
#         '·': '-',    # middle dot
#         '※': '',     # remove rare symbols
#     }
#     for k, v in replacements.items():
#         text = text.replace(k, v)
#     return text

# def extract_text_from_file(file_path, file_extension):
#     try:
#         if file_extension in ['.pdf']:
#             return extract_text_from_pdf(file_path)
#         elif file_extension in ['.docx', '.doc']:
#             return extract_text_from_docx(file_path)
#         elif file_extension in ['.txt']:
#             return extract_text_from_txt(file_path)
#         elif file_extension in ['.mp4', '.avi', '.mov', '.wmv']:
#             return extract_audio_from_video(file_path)
#         else:
#             return None
#     except Exception as e:
#         raise Exception(f"Text extraction failed: {str(e)}")

# def extract_text_from_pdf(file_path):
#     text = ""
#     with open(file_path, 'rb') as file:
#         pdf_reader = PyPDF2.PdfReader(file)
#         for page in pdf_reader.pages:
#             page_text = page.extract_text()
#             if page_text:
#                 text += page_text + "\n"
#     return text

# def extract_text_from_docx(file_path):
#     doc = docx_module.Document(file_path)
#     text = ""
#     for paragraph in doc.paragraphs:
#         text += paragraph.text + "\n"
#     return text

# def extract_text_from_txt(file_path):
#     with open(file_path, 'r', encoding='utf-8') as file:
#         return file.read()

# def extract_audio_from_video(file_path):
#     try:
#         with tempfile.NamedTemporaryFile(delete=False, suffix='.wav') as audio_tmp:
#             video = VideoFileClip(file_path)
#             video.audio.write_audiofile(audio_tmp.name, verbose=False, logger=None)
#             video.close()
#             if whisper_model is None:
#                 raise Exception("Whisper model not loaded")
#             result = whisper_model.transcribe(audio_tmp.name, word_timestamps=False)
#             os.remove(audio_tmp.name)

#             segments = result.get('segments', [])
#             transcript_lines = []
#             for seg in segments:
#                 start = seg['start']
#                 end = seg['end']
#                 text = seg['text'].strip()
#                 transcript_lines.append(f"[{start:.2f} - {end:.2f}] {text}")

#             return "\n".join(transcript_lines)
#     except Exception as e:
#         raise Exception(f"Video processing failed: {str(e)}")

# def generate_summary(text, language='en'):
#     try:
#         if not gemini_configured:
#             return "Gemini AI not configured"
            
#         language_names = {
#             'en': 'English',
#             'es': 'Spanish',
#             'fr': 'French',
#             'de': 'German',
#             'zh-cn': 'Chinese',
#             'ar': 'Arabic',
#             'hi': 'Hindi'
#         }
#         language_name = language_names.get(language, 'English')
#         model_gemini = genai.GenerativeModel('gemini-2.0-flash')
#         prompt = f"""
#         Analyze this transcript and create a comprehensive, well-structured summary in {language_name} with the following sections:

#         1. EXECUTIVE SUMMARY: A brief overview of the entire meeting
#         2. KEY POINTS: Bulleted list of the most important discussion points
#         3. ACTION ITEMS: Clear, actionable tasks with assigned responsibilities and deadlines
#         4. DECISIONS MADE: Important decisions reached during the meeting
#         5. NEXT STEPS: Planned follow-up actions and timeline

#         Format the response in a clean, professional structure with clear section headings.

#         Transcript: {text}
#         """
#         response = model_gemini.generate_content(prompt)
#         return response.text
#     except Exception as e:
#         return f"Summary generation failed: {str(e)}"

# def create_pdf(content, title, include_summary=False, summary_content=""):
#     pdf = FPDF()
#     pdf.add_page()

#     pdf.set_font("Arial", 'B', 16)
#     pdf.cell(0, 10, sanitize_text(title), 0, 1, 'C')
#     pdf.ln(10)

#     pdf.set_font("Arial", 'I', 10)
#     pdf.cell(0, 10, f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M')}", 0, 1, 'C')
#     pdf.ln(10)

#     pdf.set_font("Arial", size=12)
#     pdf.multi_cell(0, 8, sanitize_text(content))

#     if include_summary and summary_content:
#         pdf.add_page()
#         pdf.set_font("Arial", 'B', 14)
#         pdf.cell(0, 10, "AI SUMMARY", 0, 1)
#         pdf.ln(5)
#         pdf.set_font("Arial", size=12)
#         pdf.multi_cell(0, 8, sanitize_text(summary_content))

#     return pdf

# def create_docx(content, title, include_summary=False, summary_content=""):
#     doc = Document()
#     doc.add_heading(sanitize_text(title), 0)
#     doc.add_paragraph(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
#     doc.add_paragraph(sanitize_text(content))

#     if include_summary and summary_content:
#         doc.add_heading("AI Summary", 1)
#         doc.add_paragraph(sanitize_text(summary_content))

#     return doc
import whisper
from deep_translator import GoogleTranslator
import tempfile
import os
from fpdf import FPDF
from docx import Document
import google.generativeai as genai
import PyPDF2
import docx as docx_module
from moviepy.video.io.VideoFileClip import VideoFileClip
from datetime import datetime
import logging
from textblob import TextBlob  # For sentiment analysis

logger = logging.getLogger(__name__)

# Global variables for models
whisper_model = None
gemini_configured = False


def init_chatbot_models(app):
    """Initialize Whisper and Gemini models"""
    global whisper_model, gemini_configured
    try:
        whisper_model = whisper.load_model("base")
        app.logger.info("Whisper model loaded successfully")
    except Exception as e:
        app.logger.error(f"Failed to load Whisper model: {str(e)}")
        whisper_model = None

    try:
        genai.configure(api_key=app.config['GEMINI_API_KEY'])
        gemini_configured = True
        app.logger.info("Gemini configured successfully")
    except Exception as e:
        app.logger.error(f"Failed to configure Gemini: {str(e)}")
        gemini_configured = False


def analyze_sentiment(text):
    """Analyze sentiment of text and return polarity and subjectivity scores"""
    try:
        blob = TextBlob(text)
        polarity = blob.sentiment.polarity  # -1 to 1
        subjectivity = blob.sentiment.subjectivity  # 0 to 1

        if polarity > 0.1:
            sentiment = "Positive"
        elif polarity < -0.1:
            sentiment = "Negative"
        else:
            sentiment = "Neutral"

        return {
            "sentiment": sentiment,
            "polarity": round(polarity, 2),
            "subjectivity": round(subjectivity, 2)
        }
    except Exception as e:
        logger.error(f"Sentiment analysis failed: {str(e)}")
        return {"sentiment": "Unknown", "polarity": 0, "subjectivity": 0}


def sanitize_text(text):
    """Replace problematic Unicode characters with ASCII equivalents for FPDF/DOCX."""
    if not text:
        return ""
    replacements = {
        '‘': "'", '’': "'",
        '“': '"', '”': '"',
        '—': '-', '–': '-',
        '…': '...',
        '•': '-',  # bullet
        '·': '-',  # middle dot
        '※': '',   # rare symbols
    }
    for k, v in replacements.items():
        text = text.replace(k, v)
    return text


def extract_text_from_file(file_path, file_extension):
    try:
        if file_extension in ['.pdf']:
            return extract_text_from_pdf(file_path)
        elif file_extension in ['.docx', '.doc']:
            return extract_text_from_docx(file_path)
        elif file_extension in ['.txt']:
            return extract_text_from_txt(file_path)
        elif file_extension in ['.mp4', '.avi', '.mov', '.wmv']:
            return extract_audio_from_video(file_path)
        else:
            return None
    except Exception as e:
        raise Exception(f"Text extraction failed: {str(e)}")


def extract_text_from_pdf(file_path):
    text = ""
    with open(file_path, 'rb') as file:
        pdf_reader = PyPDF2.PdfReader(file)
        for page in pdf_reader.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
    return text


def extract_text_from_docx(file_path):
    doc = docx_module.Document(file_path)
    text = ""
    for paragraph in doc.paragraphs:
        text += paragraph.text + "\n"
    return text


def extract_text_from_txt(file_path):
    with open(file_path, 'r', encoding='utf-8') as file:
        return file.read()


def extract_audio_from_video(file_path):
    """Extract and transcribe audio from video using Whisper"""
    try:
        with tempfile.NamedTemporaryFile(delete=False, suffix='.wav') as audio_tmp:
            video = VideoFileClip(file_path)
            video.audio.write_audiofile(audio_tmp.name, verbose=False, logger=None)
            video.close()

            if whisper_model is None:
                raise Exception("Whisper model not loaded")

            result = whisper_model.transcribe(audio_tmp.name, word_timestamps=False)
            os.remove(audio_tmp.name)

            segments = result.get('segments', [])
            transcript_lines = []
            for seg in segments:
                start = seg['start']
                end = seg['end']
                text = seg['text'].strip()
                transcript_lines.append(f"[{start:.2f} - {end:.2f}] {text}")

            return "\n".join(transcript_lines)
    except Exception as e:
        raise Exception(f"Video processing failed: {str(e)}")


def generate_summary(text, language='en'):
    """Generate structured meeting summary with sentiment analysis included"""
    try:
        if not gemini_configured:
            return "Gemini AI not configured"

        language_names = {
            'en': 'English',
            'es': 'Spanish',
            'fr': 'French',
            'de': 'German',
            'zh-cn': 'Chinese',
            'ar': 'Arabic',
            'hi': 'Hindi'
        }
        language_name = language_names.get(language, 'English')
        model_gemini = genai.GenerativeModel('gemini-2.0-flash')

        # Sentiment analysis
        sentiment = analyze_sentiment(text)

        # Prompt
        prompt = f"""
        Analyze this transcript and create a comprehensive, well-structured summary in {language_name} with the following sections:

        1. EXECUTIVE SUMMARY: A brief overview of the entire meeting
        2. KEY POINTS: Bulleted list of the most important discussion points
        3. ACTION ITEMS: Clear, actionable tasks with assigned responsibilities and deadlines
        4. DECISIONS MADE: Important decisions reached during the meeting
        5. NEXT STEPS: Planned follow-up actions and timeline
        6. SENTIMENT ANALYSIS: Overall sentiment of the meeting (positive, negative, neutral) and key points that drove that sentiment.

        The sentiment analysis has detected: {sentiment['sentiment']} 
        (Polarity: {sentiment['polarity']}, Subjectivity: {sentiment['subjectivity']}).

        Format the response in a clean, professional structure with clear section headings.

        Transcript: {text}
        """

        response = model_gemini.generate_content(prompt)
        return response.text if hasattr(response, "text") else str(response)
    except Exception as e:
        logger.error(f"Summary generation failed: {str(e)}")
        return f"Summary generation failed: {str(e)}"


def create_pdf(content, title, include_summary=False, summary_content=""):
    pdf = FPDF()
    pdf.add_page()

    # Title
    pdf.set_font("Arial", 'B', 16)
    pdf.cell(0, 10, sanitize_text(title), 0, 1, 'C')
    pdf.ln(10)

    # Date
    pdf.set_font("Arial", 'I', 10)
    pdf.cell(0, 10, f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M')}", 0, 1, 'C')
    pdf.ln(10)

    pdf.set_font("Arial", size=12)

    # Main content
    if "1. EXECUTIVE SUMMARY:" in content:
        sections = content.split("\n\n")
        for section in sections:
            if ":" in section and not section.startswith(" "):
                pdf.set_font("Arial", 'B', 14)
                pdf.cell(0, 10, sanitize_text(section.strip()), 0, 1)
                pdf.ln(5)
                pdf.set_font("Arial", size=12)
            else:
                pdf.multi_cell(0, 8, sanitize_text(section))
                pdf.ln(5)
    else:
        pdf.multi_cell(0, 8, sanitize_text(content))

    # AI summary page
    if include_summary and summary_content:
        pdf.add_page()
        pdf.set_font("Arial", 'B', 16)
        pdf.cell(0, 10, "AI SUMMARY", 0, 1, 'C')
        pdf.ln(10)

        if "1. EXECUTIVE SUMMARY:" in summary_content:
            sections = summary_content.split("\n\n")
            for section in sections:
                if ":" in section and not section.startswith(" "):
                    pdf.set_font("Arial", 'B', 14)
                    pdf.cell(0, 10, sanitize_text(section.strip()), 0, 1)
                    pdf.ln(5)
                    pdf.set_font("Arial", size=12)
                else:
                    pdf.multi_cell(0, 8, sanitize_text(section))
                    pdf.ln(5)
        else:
            pdf.set_font("Arial", size=12)
            pdf.multi_cell(0, 8, sanitize_text(summary_content))

    return pdf


def create_docx(content, title, include_summary=False, summary_content=""):
    doc = Document()

    # Title
    doc.add_heading(sanitize_text(title), 0)

    # Date
    doc.add_paragraph(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M')}")

    # Main content
    if "1. EXECUTIVE SUMMARY:" in content:
        sections = content.split("\n\n")
        for section in sections:
            if ":" in section and not section.startswith(" "):
                doc.add_heading(section.strip(), level=2)
            else:
                doc.add_paragraph(sanitize_text(section))
    else:
        doc.add_paragraph(sanitize_text(content))

    # AI summary
    if include_summary and summary_content:
        doc.add_heading("AI Summary", 1)
        if "1. EXECUTIVE SUMMARY:" in summary_content:
            sections = summary_content.split("\n\n")
            for section in sections:
                if ":" in section and not section.startswith(" "):
                    doc.add_heading(section.strip(), level=2)
                else:
                    doc.add_paragraph(sanitize_text(section))
        else:
            doc.add_paragraph(sanitize_text(summary_content))

    return doc
